"""
DOCX Parser - Extracts text from Word documents
"""
from pathlib import Path
from typing import Dict, Any
import logging

logger = logging.getLogger(__name__)


def parse_docx(file_path: Path) -> Dict[str, Any]:
    """
    Parse a DOCX file and extract text content.
    
    Returns:
        {
            'text': str,           # Full extracted text
            'paragraphs': List[str], # Text per paragraph
            'metadata': dict,      # Document metadata
            'source': str          # Original filename
        }
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("Please install python-docx: pip install python-docx")
    
    paragraphs = []
    tables_text = []
    
    try:
        doc = Document(file_path)
        
        # Extract core properties as metadata
        metadata = {}
        if doc.core_properties:
            props = doc.core_properties
            metadata = {
                'author': props.author,
                'created': str(props.created) if props.created else None,
                'modified': str(props.modified) if props.modified else None,
                'title': props.title,
                'subject': props.subject,
            }
        
        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # Extract tables
        for table in doc.tables:
            table_text = _table_to_text(table)
            if table_text:
                tables_text.append(table_text)
        
        # Combine all content
        all_content = paragraphs + tables_text
        full_text = "\n\n".join(all_content)
        
        return {
            'text': full_text,
            'paragraphs': paragraphs,
            'tables': tables_text,
            'metadata': metadata,
            'source': file_path.name
        }
        
    except Exception as e:
        logger.error(f"Error parsing DOCX {file_path}: {e}")
        return {
            'text': '',
            'paragraphs': [],
            'tables': [],
            'metadata': {},
            'source': file_path.name,
            'error': str(e)
        }


def _table_to_text(table) -> str:
    """Convert a DOCX table to readable text"""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if any(cells):  # Skip empty rows
            rows.append(" | ".join(cells))
    return "\n".join(rows)
